import os
import re
import pdfplumber
from docx import Document
import pandas as pd

class CVDomainDetector:
    def __init__(self):
        self.cv_folder = "downloaded_cvs"
        self.csv_file = "data/cv_applications.csv"
        self.processed_csv = "data/cv_with_domains.csv"
        
        self.domain_keywords = {
            "Graphic Designing": {
                "primary": ["graphic design", "graphic designer", "visual design", "brand identity", "logo design"],
                "tools": ["photoshop", "illustrator", "indesign", "coreldraw", "adobe creative suite", "figma"],
                "skills": ["typography", "layout", "branding", "illustration", "color theory", "vector graphics"],
                "experience": ["poster design", "brochure", "packaging design", "print design", "digital design"]
            },
            
            "AI Automation": {
                "primary": ["artificial intelligence", "machine learning", "automation", "ai engineer", "ml engineer"],
                "tools": ["tensorflow", "pytorch", "opencv", "pandas", "numpy", "scikit-learn", "keras"],
                "skills": ["deep learning", "neural networks", "nlp", "computer vision", "data science"],
                "experience": ["chatbot", "rpa", "workflow automation", "model training", "algorithm development"]
            },
            
            "Accounting": {
                "primary": ["accounting", "accountant", "financial analyst", "bookkeeper", "auditor"],
                "tools": ["quickbooks", "excel", "sap", "tally", "sage", "peachtree"],
                "skills": ["financial reporting", "tax preparation", "budgeting", "cost accounting", "audit"],
                "experience": ["accounts payable", "accounts receivable", "payroll", "balance sheet", "income statement"]
            },
            
            "Web Development": {
                "primary": ["web development", "web developer", "frontend developer", "backend developer"],
                "tools": ["html", "css", "javascript", "php", "mysql", "bootstrap", "jquery"],
                "skills": ["responsive design", "dom manipulation", "ajax", "rest api", "database design"],
                "experience": ["website development", "web application", "e-commerce", "cms development"]
            },
            
            "MERN Stack": {
                "primary": ["mern stack", "mern developer", "full stack javascript", "react developer", "node developer"],
                "tools": ["mongodb", "express", "react", "nodejs", "redux", "mongoose", "npm", "yarn"],
                "skills": ["jsx", "hooks", "state management", "api integration", "component development"],
                "experience": ["spa development", "real-time applications", "microservices", "full stack projects"]
            },
            
            "Full Stack Development": {
                "primary": ["full stack", "fullstack developer", "software developer", "application developer"],
                "tools": ["docker", "kubernetes", "aws", "git", "jenkins", "linux", "nginx"],
                "skills": ["microservices", "devops", "cloud computing", "api development", "database design"],
                "experience": ["end-to-end development", "system architecture", "deployment", "scalable applications"]
            },
            
            "UI/UX Design": {
                "primary": ["ui design", "ux design", "user experience", "user interface", "interaction design"],
                "tools": ["figma", "adobe xd", "sketch", "invision", "principle", "framer"],
                "skills": ["wireframing", "prototyping", "user research", "usability testing", "design thinking"],
                "experience": ["mobile app design", "web design", "user journey", "information architecture"]
            }
        }
        
        self.category_weights = {
            "primary": 15,
            "tools": 8,
            "skills": 10,
            "experience": 12
        }
    
    def read_pdf_cv(self, file_path):
        """Extract text from PDF CV"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " ".join([cell for cell in row if cell]) + "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\n+', '\n', text)
            
            return text.strip()
            
        except Exception as error:
            print(f"❌ Error reading PDF {file_path}: {error}")
            return ""
    
    def read_docx_cv(self, file_path):
        """Extract text from DOCX CV"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " ".join(row_text) + "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\n+', '\n', text)
            
            return text.strip()
            
        except Exception as error:
            print(f"❌ Error reading DOCX {file_path}: {error}")
            return ""
    
    def extract_cv_text(self, file_path):
        """Extract text from CV based on file extension"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.read_pdf_cv(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.read_docx_cv(file_path)
        else:
            print(f"⚠️ Unsupported file format: {file_extension}")
            return ""
    
    def calculate_domain_score(self, cv_text, domain_keywords):
        """Calculate domain score"""
        if not cv_text:
            return 0, []
        
        cv_text_lower = cv_text.lower()
        total_score = 0
        found_keywords = []
        
        for category, keywords in domain_keywords.items():
            category_weight = self.category_weights.get(category, 5)
            
            for keyword in keywords:
                keyword_lower = keyword.lower()
                exact_matches = cv_text_lower.count(keyword_lower)
                
                if exact_matches > 0:
                    keyword_importance = len(keyword.split()) * 2
                    frequency_bonus = min(exact_matches, 3)
                    
                    score = category_weight * keyword_importance * frequency_bonus
                    total_score += score
                    
                    found_keywords.append(keyword)
        
        return total_score, found_keywords
    
    def detect_domain(self, cv_text):
        """Detect domain with confidence scoring"""
        if not cv_text:
            return "Unknown", 0, []
        
        domain_scores = {}
        all_keywords_found = {}
        
        for domain, keywords in self.domain_keywords.items():
            score, found_keywords = self.calculate_domain_score(cv_text, keywords)
            domain_scores[domain] = score
            all_keywords_found[domain] = found_keywords
        
        max_score = max(domain_scores.values()) if domain_scores else 0
        
        if max_score == 0:
            return "Unknown", 0, []
        
        best_domain = max(domain_scores, key=domain_scores.get)
        
        max_possible_score = 500
        confidence = min(int((max_score / max_possible_score) * 100), 100)
        
        if confidence < 15:
            return "Unknown", confidence, []
        
        return best_domain, confidence, all_keywords_found[best_domain]
    
    def find_cv_file_precise(self, name, email):
        """Find CV file with precise matching"""
        if not os.path.exists(self.cv_folder):
            return None
        
        clean_name = re.sub(r'[^\w\s-]', '', name).strip().lower()
        email_prefix = email.split('@')[0].lower() if email else ""
        
        cv_files = []
        for filename in os.listdir(self.cv_folder):
            if filename.lower().endswith(('.pdf', '.docx', '.doc')):
                cv_files.append(filename)
        
        for filename in cv_files:
            filename_lower = filename.lower()
            
            if clean_name and clean_name in filename_lower:
                return os.path.join(self.cv_folder, filename)
            
            if email_prefix and len(email_prefix) > 3 and email_prefix in filename_lower:
                return os.path.join(self.cv_folder, filename)
        
        name_words = clean_name.split()
        for filename in cv_files:
            filename_lower = filename.lower()
            matches = sum(1 for word in name_words if len(word) > 2 and word in filename_lower)
            
            if matches >= len(name_words) // 2 and matches > 0:
                return os.path.join(self.cv_folder, filename)
        
        return None
    
    def process_cvs_with_domain_detection(self):
        """Main function to process all CVs and detect domains"""
        print("🚀 Starting CV processing and domain detection...")
        
        if not os.path.exists(self.csv_file):
            print(f"❌ {self.csv_file} not found. Please run Gmail scanner first.")
            return
        
        df = pd.read_csv(self.csv_file)
        print(f"📊 Found {len(df)} CV records to process")
        
        if 'domain' not in df.columns:
            df['domain'] = ""
        if 'confidence' not in df.columns:
            df['confidence'] = 0
        if 'keywords_found' not in df.columns:
            df['keywords_found'] = ""
        if 'cv_text_preview' not in df.columns:
            df['cv_text_preview'] = ""
        
        processed_count = 0
        
        for index, row in df.iterrows():
            name = row['name']
            email = row['email']
            
            print(f"\n📄 Processing CV for: {name} ({email})")
            
            cv_file = self.find_cv_file_precise(name, email)
            
            if not cv_file:
                print(f"❌ CV file not found for {name}")
                df.at[index, 'domain'] = "File Not Found"
                df.at[index, 'confidence'] = 0
                df.at[index, 'keywords_found'] = ""
                df.at[index, 'cv_text_preview'] = ""
                continue
            
            cv_text = self.extract_cv_text(cv_file)
            
            if not cv_text or len(cv_text.strip()) < 50:
                print(f"❌ Could not extract sufficient text from {cv_file}")
                df.at[index, 'domain'] = "Text Extraction Failed"
                df.at[index, 'confidence'] = 0
                df.at[index, 'keywords_found'] = ""
                df.at[index, 'cv_text_preview'] = "Text extraction failed"
                continue
            
            cv_preview = cv_text[:1000].strip()
            df.at[index, 'cv_text_preview'] = cv_preview
            
            domain, confidence, keywords_found = self.detect_domain(cv_text)
            
            df.at[index, 'domain'] = domain
            df.at[index, 'confidence'] = confidence
            df.at[index, 'keywords_found'] = ", ".join(keywords_found) if keywords_found else ""
            
            processed_count += 1
            print(f"✅ {name} → {domain} | Confidence: {confidence}%")
            
            if processed_count % 5 == 0:
                df.to_csv(self.processed_csv, index=False)
                print(f"💾 Saved progress: {processed_count} candidates processed")
        
        df.to_csv(self.processed_csv, index=False)
        print(f"\n🎉 Processing complete!")
        print(f"📊 Processed {processed_count} CVs")
        print(f"💾 Results saved to: {self.processed_csv}")
        
        self.show_domain_summary(df)
        
        return df
    
    def show_domain_summary(self, df):
        """Show summary of domain distribution"""
        print("\n📈 Domain Detection Summary:")
        print("-" * 50)
        
        domain_counts = df['domain'].value_counts()
        
        for domain, count in domain_counts.items():
            avg_confidence = df[df['domain'] == domain]['confidence'].mean()
            print(f"{domain}: {count} candidates (avg confidence: {avg_confidence:.1f}%)")
        
        print(f"\nTotal: {len(df)} candidates processed")
        
        high_confidence = df[df['confidence'] >= 50]
        print(f"High confidence detections (≥50%): {len(high_confidence)}")

if __name__ == "__main__":
    detector = CVDomainDetector()
    
    print("Choose processing method:")
    print("1. Process CVs from Gmail scanner data")
    print("2. Show domain summary")
    
    choice = input("Enter your choice (1-2): ")
    
    if choice == "1":
        detector.process_cvs_with_domain_detection()
    elif choice == "2":
        if os.path.exists(detector.processed_csv):
            df = pd.read_csv(detector.processed_csv)
            detector.show_domain_summary(df)
        else:
            print("❌ No processed data found")
    else:
        print("Invalid choice. Running default method...")
        detector.process_cvs_with_domain_detection()